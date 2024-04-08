import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import unidecode
import csv
import os
import docx

data = docx.Document("0.docx")

for k in range(5, 7):
    example = []
    tmp = []
    for i in range(len(data.tables[k].rows)):
        tmp.clear()
        for j in range(len(data.tables[k].rows[i].cells)):
            tmp.append(data.tables[k].rows[i].cells[j].text)

        example.append(np.array(tmp))

    if k == 9:
        example.pop(32)
        example.pop(31)
    example = np.array(example)
    result = pd.DataFrame(example)
    result.to_excel("" + str(k) + ".xlsx")
