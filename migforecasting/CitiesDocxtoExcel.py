import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import unidecode
import csv
import os
import docx

data = docx.Document("0.docx")

example = []
tmp = []
n = 4
x = 1
if len(data.tables) == 7 + x:    # данные разбиты на 4 таблицы
    a = data.tables[n]
    b = data.tables[n + 1]
    c = data.tables[n + 2]
    d = data.tables[n + 3]

    for i in range(len(a.rows)):
        tmp.clear()
        for j in range(len(a.rows[i].cells)):
            tmp.append(a.rows[i].cells[j].text)

        for j in range(len(b.rows[i].cells)):
            tmp.append(b.rows[i].cells[j].text)

        example.append(np.array(tmp))

    for i in range(2, len(c.rows)):
        tmp.clear()
        for j in range(len(c.rows[i].cells)):
            tmp.append(c.rows[i].cells[j].text)

        for j in range(len(d.rows[i].cells)):
            tmp.append(d.rows[i].cells[j].text)

        example.append(np.array(tmp))
elif len(data.tables) == 5 + x:   # данные разбиты на 2 таблицы
    a = data.tables[n]
    b = data.tables[n + 1]
    for i in range(len(a.rows)):
        tmp.clear()

        for j in range(len(a.rows[i].cells)):
            tmp.append(a.rows[i].cells[j].text)

        w = tmp[0].split()
        if len(w) != 0:
            if w[0] == 'Миграционный' and n == 4:
                for y in range(5):
                    empt = list(range(len(a.rows[i].cells)))
                    example.append(np.array(empt))

        try:
            if w[0] + w[1] == 'Численностьдетей,' and n == 4:
                tmp.clear()
            else:
                example.append(np.array(tmp))
        except IndexError:
            example.append(np.array(tmp))


    for i in range(1, len(b.rows)):
        tmp.clear()
        if len(b.rows[i].cells) == len(a.rows[0].cells):
            for j in range(len(b.rows[i].cells)):
                tmp.append(b.rows[i].cells[j].text)
        else:
            tmp.append(b.rows[i].cells[0].text)
            tmp.append(b.rows[i].cells[2].text)
            tmp.append(b.rows[i].cells[5].text)
            tmp.append(b.rows[i].cells[6].text)

        example.append(np.array(tmp))

elif len(data.tables) == 6 + x:     # данные разбиты на 3 таблицы
    v = 2
    for k in range(n, 6):
        if k == 3:
            v = 0
        else:
            v = 2
        for i in range(v, len(data.tables[k].rows)):
            tmp.clear()
            for j in range(len(data.tables[k].rows[i].cells)):
                tmp.append(data.tables[k].rows[i].cells[j].text)

            example.append(np.array(tmp))


example = np.array(example)

result = pd.DataFrame(example)

result.to_excel("d0.xlsx")